import os
import re
from urllib.parse import quote
from markdown_pdf import MarkdownPdf, Section
from dotenv import load_dotenv

load_dotenv()

# 项目根路径（无论你在哪里运行）
BASE_DIR = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# 从 .env 获取 DATA_DIR，相对于 BASE_DIR 解析
DATA_DIR_NAME = os.getenv("DATA_DIR", "data")
DATA_DIR = os.path.join(BASE_DIR, DATA_DIR_NAME)
SAVE_PATH = os.path.join(DATA_DIR, "note_output")
IMAGE_BASE_URL = os.getenv("IMAGE_BASE_URL")
STATIC_BASE = os.path.join(BASE_DIR, IMAGE_BASE_URL)


class ExportUtils:
    def __init__(self, **kwargs):
        # 确认SAVE_PATH存在
        print(f"保存路径: {SAVE_PATH}")
        print(f"静态文件路径: {STATIC_BASE}")
        if not os.path.exists(SAVE_PATH):
            os.makedirs(SAVE_PATH)

    def _embed_image_as_base64(self, img_path: str) -> str:
        """
        将图片转换为 base64 格式嵌入
        """
        import base64
        import mimetypes

        try:
            # 获取 MIME 类型
            mime_type, _ = mimetypes.guess_type(img_path)
            if not mime_type:
                # 根据扩展名推断
                ext = os.path.splitext(img_path)[1].lower()
                mime_map = {
                    '.png': 'image/png',
                    '.jpg': 'image/jpeg',
                    '.jpeg': 'image/jpeg',
                    '.gif': 'image/gif',
                    '.bmp': 'image/bmp',
                    '.webp': 'image/webp',
                    '.svg': 'image/svg+xml'
                }
                mime_type = mime_map.get(ext, 'image/png')

            # 读取图片文件并转换为 base64
            with open(img_path, 'rb') as f:
                img_data = f.read()

            base64_data = base64.b64encode(img_data).decode('utf-8')
            return f"data:{mime_type};base64,{base64_data}"

        except Exception as e:
            print(f"图片 base64 编码失败 {img_path}: {str(e)}")
            return None

    def _get_normalized_path(self, path: str) -> str:
        """
        获取规范化的绝对路径
        """
        return os.path.normpath(os.path.abspath(path))

    def _replace_static_paths_with_absolute(self, content: str) -> str:
        """
        将 Markdown 中的图片路径替换为 base64 内嵌格式
        这样可以确保图片在 PDF 中正确显示
        """

        def repl(match):
            # 捕获 alt 文本和路径
            alt_text = match.group(1) if match.group(1) else ""
            img_path = match.group(2).strip()

            print(f"处理图片路径: {img_path}")

            # 处理 /static/ 开头的路径
            if img_path.startswith("/static/"):
                # 构建绝对路径
                relative_path = img_path.lstrip("/")  # 移除开头的 /
                abs_path = os.path.join(BASE_DIR, relative_path)
                abs_path = self._get_normalized_path(abs_path)

                # 检查文件是否存在并转换为 base64
                if os.path.exists(abs_path):
                    base64_uri = self._embed_image_as_base64(abs_path)
                    if base64_uri:
                        print(f"图片转换为 base64 成功: {img_path}")
                        return f"![{alt_text}]({base64_uri})"
                    else:
                        print(f"图片 base64 转换失败: {abs_path}")
                        return f"![{alt_text}](图片转换失败: {img_path})"
                else:
                    print(f"警告：图片文件不存在 {abs_path}")
                    return f"![{alt_text}](图片不存在: {img_path})"

            # 处理相对路径（相对于 STATIC_BASE）
            elif not img_path.startswith(('http://', 'https://', 'data:')):
                # 尝试多个可能的路径
                possible_paths = [
                    os.path.join(STATIC_BASE, img_path),
                    os.path.abspath(img_path),
                    os.path.join(BASE_DIR, img_path)
                ]

                for abs_path in possible_paths:
                    abs_path = self._get_normalized_path(abs_path)
                    if os.path.exists(abs_path):
                        base64_uri = self._embed_image_as_base64(abs_path)
                        if base64_uri:
                            print(f"相对路径图片转换为 base64 成功: {img_path}")
                            return f"![{alt_text}]({base64_uri})"
                        break

                print(f"警告：图片文件未找到 {img_path}")
                return f"![{alt_text}](图片未找到: {img_path})"

            # HTTP/HTTPS 和 data: 路径保持不变
            elif img_path.startswith(('http://', 'https://', 'data:')):
                print(f"网络图片或 data URI 保持不变: {img_path[:50]}...")
                return match.group(0)

            # 其他情况保持不变
            return match.group(0)

        # 使用更精确的正则表达式匹配图片语法
        # 匹配 ![alt text](path) 格式
        pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
        result = re.sub(pattern, repl, content)

        print("图片路径处理完成")
        return result

    def _to_html(self, content: str, title: str):
        """
        将 Markdown 内容转换为 HTML。
        用 markdown-it-py 做转换，套一个简洁中文友好的 CSS。
        """
        try:
            from markdown_it import MarkdownIt

            md = MarkdownIt("commonmark", {"html": True, "linkify": True, "typographer": True}).enable("table").enable("strikethrough")
            html_body = md.render(content)

            css = """
            body {
              font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", "PingFang SC",
                           "Hiragino Sans GB", "Microsoft YaHei", "Source Han Sans CN",
                           Roboto, Helvetica, Arial, sans-serif;
              line-height: 1.7;
              color: #24292f;
              max-width: 860px;
              margin: 32px auto;
              padding: 0 24px 64px;
            }
            h1, h2, h3, h4 { line-height: 1.3; margin-top: 1.4em; margin-bottom: 0.6em; font-weight: 600; }
            h1 { font-size: 2em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }
            h2 { font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }
            h3 { font-size: 1.25em; }
            p { margin: 0.8em 0; }
            a { color: #0969da; text-decoration: none; }
            a:hover { text-decoration: underline; }
            code { background: rgba(175,184,193,0.2); padding: 0.18em 0.4em; border-radius: 4px;
                   font-family: "SFMono-Regular", Consolas, "Liberation Mono", Menlo, monospace; font-size: 0.92em; }
            pre { background: #f6f8fa; padding: 14px 16px; border-radius: 6px; overflow-x: auto; }
            pre code { background: none; padding: 0; }
            blockquote { border-left: 4px solid #d0d7de; color: #57606a; padding: 0 1em; margin: 1em 0; }
            img { max-width: 100%; border-radius: 6px; }
            table { border-collapse: collapse; margin: 1em 0; }
            table th, table td { border: 1px solid #d0d7de; padding: 6px 13px; }
            table th { background: #f6f8fa; font-weight: 600; }
            ul, ol { padding-left: 1.6em; }
            hr { border: 0; border-top: 1px solid #eaecef; margin: 2em 0; }
            """

            html_doc = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="utf-8" />
<meta name="viewport" content="width=device-width, initial-scale=1" />
<title>{title}</title>
<style>{css}</style>
</head>
<body>
<h1>{title}</h1>
{html_body}
</body>
</html>
"""
            save_path = os.path.join(SAVE_PATH, f"{title}.html")
            with open(save_path, "w", encoding="utf-8") as f:
                f.write(html_doc)
            print(f"HTML 导出成功: {save_path}")
            return save_path
        except Exception as e:
            print(f"HTML 导出失败: {e}")
            raise

    def _to_word(self, content: str, title: str):
        """
        将 Markdown 内容转换为 Word (.docx)。
        策略：markdown → HTML（已有 _to_html 的 md），然后用 BeautifulSoup 遍历元素塞进 python-docx。
        覆盖标题（h1-h4）、段落、有/无序列表、代码块、链接（保留文本）、图片（base64 已被前置处理，
        先解出来再插入），其它当作普通段落处理。
        """
        try:
            from io import BytesIO
            import base64 as _b64
            from markdown_it import MarkdownIt
            from bs4 import BeautifulSoup
            from docx import Document
            from docx.shared import Pt, Inches

            md = MarkdownIt("commonmark", {"html": True}).enable("table").enable("strikethrough")
            html = md.render(content)
            soup = BeautifulSoup(html, "html.parser")

            doc = Document()
            # 默认正文字体
            style = doc.styles["Normal"]
            style.font.name = "Microsoft YaHei"
            style.font.size = Pt(11)

            doc.add_heading(title, level=0)

            heading_levels = {"h1": 1, "h2": 2, "h3": 3, "h4": 4, "h5": 5, "h6": 6}

            def _add_image(src: str):
                try:
                    if src.startswith("data:"):
                        # data:image/png;base64,XXXX
                        header, _, b64 = src.partition(",")
                        img_bytes = _b64.b64decode(b64)
                        doc.add_picture(BytesIO(img_bytes), width=Inches(5.5))
                    elif src.startswith(("http://", "https://")):
                        doc.add_paragraph(f"[网络图片] {src}")
                    elif os.path.exists(src):
                        doc.add_picture(src, width=Inches(5.5))
                    else:
                        doc.add_paragraph(f"[图片缺失] {src}")
                except Exception as e:
                    doc.add_paragraph(f"[图片插入失败] {src}: {e}")

            def _add_paragraph_with_inline(el):
                """段落级元素：保留文本，链接/strong/em 都按纯文本处理（python-docx 内联格式比较繁，先够用）。"""
                text = el.get_text("", strip=False).strip()
                if text:
                    doc.add_paragraph(text)
                # 段落内嵌的图片单独提出来
                for img in el.find_all("img"):
                    _add_image(img.get("src", ""))

            for child in soup.children:
                name = getattr(child, "name", None)
                if name is None:
                    continue
                if name in heading_levels:
                    doc.add_heading(child.get_text(strip=True), level=heading_levels[name])
                elif name == "p":
                    _add_paragraph_with_inline(child)
                elif name in ("ul", "ol"):
                    style_name = "List Bullet" if name == "ul" else "List Number"
                    for li in child.find_all("li", recursive=False):
                        doc.add_paragraph(li.get_text(strip=True), style=style_name)
                elif name == "pre":
                    code_text = child.get_text("", strip=False)
                    p = doc.add_paragraph()
                    run = p.add_run(code_text)
                    run.font.name = "Consolas"
                    run.font.size = Pt(10)
                elif name == "blockquote":
                    p = doc.add_paragraph(child.get_text(strip=True))
                    p.paragraph_format.left_indent = Inches(0.3)
                elif name == "img":
                    _add_image(child.get("src", ""))
                elif name == "hr":
                    doc.add_paragraph("———————————————")
                elif name == "table":
                    rows = child.find_all("tr")
                    if not rows:
                        continue
                    cols = max(len(r.find_all(["td", "th"])) for r in rows)
                    table = doc.add_table(rows=len(rows), cols=cols)
                    table.style = "Light Grid Accent 1"
                    for r_idx, tr in enumerate(rows):
                        cells = tr.find_all(["td", "th"])
                        for c_idx, td in enumerate(cells):
                            table.cell(r_idx, c_idx).text = td.get_text(strip=True)
                else:
                    # 兜底：当成段落
                    text = child.get_text("", strip=False).strip()
                    if text:
                        doc.add_paragraph(text)

            save_path = os.path.join(SAVE_PATH, f"{title}.docx")
            doc.save(save_path)
            print(f"Word 导出成功: {save_path}")
            return save_path
        except Exception as e:
            print(f"Word 导出失败: {e}")
            raise

    def _to_pdf(self, content: str, title: str):
        """
        将 Markdown 内容转换为 PDF
        """
        try:
            # 创建 PDF 对象，启用优化
            pdf = MarkdownPdf(
                optimize=True,
                # 添加一些可能有助于图片显示的配置
                # toc=False,
                # paper_size='A4',
                # margin=dict(top='1cm', bottom='1cm', left='1cm', right='1cm')
            )

            # 添加内容段落
            pdf.add_section(Section(content))

            # 保存 PDF
            save_path = os.path.join(SAVE_PATH, f"{title}.pdf")
            pdf.save(save_path)

            print(f"PDF 导出成功: {save_path}")
            return save_path

        except Exception as e:
            print(f"PDF 导出失败: {str(e)}")
            print("尝试使用基本配置...")
            try:
                # 尝试最基本的配置
                pdf = MarkdownPdf()
                pdf.add_section(Section(content))
                save_path = os.path.join(SAVE_PATH, f"{title}.pdf")
                pdf.save(save_path)
                print(f"基本配置 PDF 导出成功: {save_path}")
                return save_path
            except Exception as e2:
                print(f"基本配置也失败: {str(e2)}")
                raise e2

    def export(self, output_format: str, title: str, content: str) -> str:
        """
        导出内容为指定格式
        支持格式：pdf, html, word/docx, image/png
        """
        content = content.strip()

        # 处理图片路径
        print("开始处理图片路径...")
        content = self._replace_static_paths_with_absolute(content)

        output_format = output_format.lower()

        try:
            if output_format == "pdf":
                save_path = self._to_pdf(content, title)
            elif output_format == "html":
                save_path = self._to_html(content, title)
            elif output_format in ["word", "docx"]:
                save_path = self._to_word(content, title)
            else:
                supported_formats = ["pdf", "html", "word/docx"]
                raise ValueError(f"不支持的导出格式: {output_format}. 支持的格式: {', '.join(supported_formats)}")

            print(f"导出完成: {save_path}")
            return save_path

        except Exception as e:
            print(f"导出失败: {str(e)}")
            raise e

    def get_supported_formats(self):
        """
        返回支持的导出格式列表
        """
        return {
            "pdf": "PDF 文档",
            "html": "HTML 网页",
            "word": "Word 文档 (.docx)",
            "docx": "Word 文档 (.docx)",
        }
    def debug_paths(self):
        """
        调试方法：打印重要路径信息
        """
        print("=== 路径调试信息 ===")
        print(f"BASE_DIR: {BASE_DIR}")
        print(f"DATA_DIR: {DATA_DIR}")
        print(f"SAVE_PATH: {SAVE_PATH}")
        print(f"STATIC_BASE: {STATIC_BASE}")
        print(f"IMAGE_BASE_URL: {IMAGE_BASE_URL}")
        print("==================")

if __name__ == '__main__':

    ExportUtils().export("pdf",title='测试',content='''# 视频笔记：Facial Recognition Forces My Coworkers to Do Their Dishes

## 简介
该视频展示了团队如何利用面部识别技术来监控和激励同事清洗餐具。通过结合硬件和软件，团队开发了一个“Dish Watcher”系统，旨在识别并提醒那些未清洁餐具的人。

## 背景
- 团队面临的问题是同事们不愿意清洗餐具。
- 为解决这一问题，团队决定在不告知的情况下使用技术来监控厨房区域。

## 实验设计
1\. **设备安装**
- 使用Raspberry Pi和隐藏摄像头来捕捉厨房水槽的活动。
- 摄像头只在有人在水槽附近活动时录制，以节省存储空间。

2\. **软件开发**
- 使用Cursor AI和Meta的项目来分析视频。
- 系统能识别人员特征如发型、服装，并将结果发送到Discord服务器以提醒团队。

3\. **面部识别**
- 通过视频流实时分析来判断是否有人留下了脏餐具。
- 系统能识别并记录下未清洗餐具的人的详细特征。

![](/static/screenshots/screenshot_000_a61be29d-06ae-42ee-ac38-2d0b1db394f3.jpg)* 展示了堆积的脏餐具，问题的严重性可见一斑。

## 实验过程
- 系统成功捕获了少数“罪犯”，并通过Discord进行了通知。
- 计划将摄像头隐藏在厨房的画作后，使其更加隐蔽。

![](/static/screenshots/screenshot_001_e9d1c7ad-509e-4c7d-a718-a09193e97724.jpg)* SAM 介绍了项目的背景。

## 结果
- 实验初期，系统有效地识别了不清洗餐具的同事。
- 由于摄像头的存在，同事们开始自觉清洗餐具，长时间未发现新的“罪犯”。

## 思考与改进
- 团队意识到仅仅通过惩罚来改变行为可能效果有限，考虑奖励来激励清洗餐具。
- 系统将改进为奖励机制，记录并表扬那些清洗餐具的人。

## 总结
这次实验展示了技术在工作场所行为管理中的应用潜力。通过实验，团队不仅解决了餐具清洗的问题，还对如何更有效地激励员工有了更深的认识。

![](/static/screenshots/screenshot_002_f1ca0c20-c657-417f-be78-7958bf0e7a4b.jpg)* 展示了系统对某位同事洗碗的实时面部识别。

## 结论
- 应用技术可以有效改善工作环境中的小问题。
- 积极的激励比惩罚更能驱动行为改变。

通过这次实验，团队不仅解决了餐具堆积的问题，还为未来更复杂的行为管理系统奠定了基础。 ''',)

